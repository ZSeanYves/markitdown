#!/usr/bin/env python3
"""Generate a minimal DOCX sample with one inline hyperlink.

Output (from repo root):
  samples/docx/docx_hyperlink_basic.docx

Dependencies:
  pip install python-docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


REPO_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = REPO_ROOT / "samples" / "docx"
OUT_FILE = OUT_DIR / "docx_hyperlink_basic.docx"


def add_hyperlink(paragraph, text: str, url: str) -> None:
    """Append a hyperlink run to an existing python-docx paragraph.

    python-docx has no public add_hyperlink API, so we create the underlying
    OOXML nodes directly:
      <w:hyperlink r:id="...">
        <w:r><w:rPr><w:rStyle w:val="Hyperlink"/></w:rPr><w:t>text</w:t></w:r>
      </w:hyperlink>
    """

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)

    text_elem = OxmlElement("w:t")
    text_elem.text = text

    new_run.append(r_pr)
    new_run.append(text_elem)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_paragraph("This is a normal paragraph before the hyperlink.")

    p = doc.add_paragraph("Visit ")
    add_hyperlink(p, text="OpenAI", url="https://openai.com")
    p.add_run(" for more information.")

    doc.save(OUT_FILE)
    print(f"[ok] wrote: {OUT_FILE}")


if __name__ == "__main__":
    main()
